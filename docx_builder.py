# docx_builder.py — structured DOCX with Immediate Actions & Advice sections
# returns bytes for Streamlit download
import io
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- helpers ---------------------------------------------------------------
def _shade(cell, hexcolor: str):
    """Fill a cell background with hexcolor (e.g., 'FDECEA')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hexcolor)
    tc_pr.append(shd)

def _h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)

def _p(doc: Document, text: str = "", bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p

def _risk_colors(risk: str):
    # background, border-ish text color
    m = {
        "HIGH":  ("FDECEA", "D93025"),  # light red, red text
        "MEDIUM":("FFF4E5", "B93815"),  # light orange, dark orange text
        "LOW":   ("E8F5E9", "1E7D32"),  # light green, green text
    }
    return m.get(risk, ("F3F4F6", "111827"))

def _domain_items(dom: Dict) -> List[Dict]:
    return [it for it in (dom.get("items") or [])]

def _as_issue_text(item: Dict) -> str:
    """Prefer a concise risk comment; fall back to Q(A) summary."""
    if item.get("risk_comment"):
        return str(item["risk_comment"])
    q = str(item.get("question", "")).strip()
    a = str(item.get("answer", "")).strip()
    if q and a:
        return f"{q}  —  Answer: {a}"
    return q or a or "Issue"

def _next_step_text(item: Dict) -> str:
    return item.get("next_step") or "Add recommended action…"

def _legal_text(item: Dict) -> str:
    # Accept either 'legal_basis' (string) or 'cases' (list)
    if item.get("legal_basis"):
        return f"Legal notes: {item['legal_basis']}"
    cases = item.get("cases") or item.get("case_law") or []
    if isinstance(cases, list) and cases:
        return "Relevant case law: " + "; ".join(map(str, cases))
    return "Relevant case law: (to be added)"

# --- core -------------------------------------------------------------------
def build_docx_bytes(data: dict) -> bytes:
    doc = Document()

    # ---------- cover ----------
    title = doc.add_paragraph()
    run = title.add_run(data.get("organisation", "Competition Risk Diagnostic"))
    run.font.size = Pt(22); run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _p(doc, f'Completed by: {data.get("completed_by","")}')
    _p(doc, f'Generated: {data.get("generated_at","")}')
    _p(doc, f'Overall risk: {data.get("overall_risk","")}')
    doc.add_paragraph()

    # ---------- domain summary table ----------
    _h(doc, "Domain summary", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Domain", "Score", "Risk"

    high_domains, med_domains, low_domains = [], [], []
    for d in data.get("domains", []):
        row = tbl.add_row().cells
        row[0].text, row[1].text, row[2].text = d["name"], str(int(d.get("score", 0))), d.get("risk", "")
        bg, tc = _risk_colors(d.get("risk", ""))
        _shade(row[2], bg)  # color the risk cell
        if d.get("risk") == "HIGH":
            high_domains.append(d)
        elif d.get("risk") == "MEDIUM":
            med_domains.append(d)
        else:
            low_domains.append(d)
    doc.add_paragraph()

    # ---------- Immediate actions (HIGH risk) ----------
    _h(doc, "Immediate actions (HIGH risk)", level=1)
    if not high_domains:
        _p(doc, "No domains scored HIGH. Review medium-risk areas below.")
    for d in high_domains:
        bg, tc = _risk_colors("HIGH")
        _h(doc, f'{d["name"]}', level=2)
        # Action table
        t = doc.add_table(rows=1, cols=3)
        t.allow_autofit = True
        t.columns[0].width = Inches(3.6)
        t.columns[1].width = Inches(2.7)
        t.columns[2].width = Inches(1.6)
        hh = t.rows[0].cells
        hh[0].text, hh[1].text, hh[2].text = "Issue", "Recommended next step", "Owner / notes"
        _shade(hh[0], bg); _shade(hh[1], bg); _shade(hh[2], bg)

        items = [it for it in _domain_items(d) if int(it.get("points", 0)) >= 0]  # include zeros if curated via next_step
        if not items:
            r = t.add_row().cells
            r[0].text = "Review this domain — high aggregate risk."
            r[1].text = "Identify and prioritise corrective actions."
            r[2].text = ""
        else:
            for it in items:
                r = t.add_row().cells
                r[0].text = _as_issue_text(it)
                r[1].text = _next_step_text(it)
                r[2].text = ""  # left blank for assignee/notes
                # Optional legal notes as a small paragraph under the row
            doc.add_paragraph(_legal_text(items[0]))  # 1 legal line under the table (keeps layout neat)
        doc.add_paragraph()

    # ---------- Advice recommended (MEDIUM risk) ----------
    _h(doc, "Advice recommended (MEDIUM risk)", level=1)
    if not med_domains:
        _p(doc, "No domains scored MEDIUM.")
    for d in med_domains:
        _h(doc, d["name"], level=2)
        _p(doc, "We recommend seeking specialist competition law advice on the points below.")
        t = doc.add_table(rows=1, cols=3)
        t.columns[0].width = Inches(3.6)
        t.columns[1].width = Inches(2.7)
        t.columns[2].width = Inches(1.6)
        hh = t.rows[0].cells
        hh[0].text, hh[1].text, hh[2].text = "Topic / concern", "Why it matters / next step", "Owner / notes"
        items = [it for it in _domain_items(d) if int(it.get("points", 0)) >= 0]
        if not items:
            r = t.add_row().cells
            r[0].text = "General review"
            r[1].text = "Clarify rules, document decision-making, and obtain advice on grey areas."
            r[2].text = ""
        else:
            for it in items:
                r = t.add_row().cells
                r[0].text = _as_issue_text(it)
                why = it.get("risk_comment") or "Potential competition sensitivity. Confirm approach and guardrails."
                r[1].text = it.get("next_step") or why
                r[2].text = ""
        doc.add_paragraph()

    # ---------- Notes & additional explanation (for the compliance officer) ----------
    _h(doc, "Additional explanation & follow-up notes", level=1)
    _p(doc, "Use this section to add context, evidence, owners and dates for each domain.")
    for d in data.get("domains", []):
        _h(doc, d["name"], level=2)
        t = doc.add_table(rows=2, cols=3)
        t.columns[0].width = Inches(3.6)
        t.columns[1].width = Inches(2.3)
        t.columns[2].width = Inches(1.6)
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Context / explanation", "Evidence / documents", "Owner / due date"
        # blank row for user to fill
        r = t.rows[1].cells
        r[0].text, r[1].text, r[2].text = "", "", ""
        doc.add_paragraph()

    # (No Q&A appendix; the report is concise and action-oriented)

    # ---------- output ----------
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
